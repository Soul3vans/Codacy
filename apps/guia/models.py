"""Modelo mejorado para Guías de Autocontrol con optimizaciones de rendimiento y API de datos"""
from django.db import models
from django.contrib.auth import get_user_model
from django.urls import reverse
from django.core.cache import cache
from django.db.models import Index
from docx import Document
from apps.dashboard.models import Archivo
import re
import logging
import PyPDF2
import mimetypes
import hashlib
from django.core.validators import MinValueValidator, MaxValueValidator

User = get_user_model()
logger = logging.getLogger(__name__)

class GuiaAutocontrol(models.Model):
    """
    Modelo mejorado para Guías de Autocontrol con:
    - Optimización de rendimiento mediante caching
    - Mejor API de datos para consultas frecuentes
    - Indexación de búsquedas
    """
    archivo = models.OneToOneField(
        Archivo,
        on_delete=models.CASCADE,
        related_name='guia_autocontrol',
        limit_choices_to={'es_formulario': True}
    )
    titulo_guia = models.CharField(max_length=300, blank=True)
    componente = models.CharField(max_length=200, blank=True)
    proposito = models.TextField(blank=True)
    contenido_procesado = models.JSONField(default=dict, blank=True)
    activa = models.BooleanField(default=True)
    fecha_procesamiento = models.DateTimeField(auto_now=True)
    hash_archivo = models.CharField(max_length=64, blank=True, editable=False)
    version = models.CharField(max_length=20, blank=True)
    
    # Campos denormalizados para optimización
    total_preguntas = models.PositiveIntegerField(default=0, editable=False)
    categorias_count = models.PositiveIntegerField(default=0, editable=False)

    class Meta:
        verbose_name = 'Guía de Autocontrol'
        verbose_name_plural = 'Guías de Autocontrol'
        ordering = ['-fecha_procesamiento']
        indexes = [
            Index(fields=['componente']),
            Index(fields=['activa', 'fecha_procesamiento']),
            Index(fields=['total_preguntas']),
        ]

    def __str__(self):
        return f"Guía: {self.titulo_guia or self.archivo.nombre}"

    def save(self, *args, **kwargs):
        """Sobreescritura de save para calcular campos denormalizados y hash"""
        if self.archivo and self.archivo.archivo:
            self.hash_archivo = self.calcular_hash_archivo()
        
        # Actualizar campos denormalizados
        if self.contenido_procesado:
            self.total_preguntas = self._calcular_total_preguntas()
            self.categorias_count = len(self.contenido_procesado.get('tablas_cuestionario', []))
        # Actualiza el título de la guía si el archivo tiene nombre
        if self.archivo and self.archivo.get_nombre_archivo():
            self.titulo_guia = self.archivo.get_nombre_archivo()
            
        super().save(*args, **kwargs)

    def calcular_hash_archivo(self):
        """Calcula hash SHA-256 del archivo para detectar cambios"""
        hash_obj = hashlib.sha256()
        with self.archivo.archivo.open('rb') as f:
            for chunk in iter(lambda: f.read(4096), b''):
                hash_obj.update(chunk)
        return hash_obj.hexdigest()

    def _calcular_total_preguntas(self):
        """Calcula el total de preguntas para denormalización"""
        count = 0
        for categoria in self.contenido_procesado.get('tablas_cuestionario', []):
            for bloque in categoria.get('bloques', []):
                count += len(bloque.get('preguntas', []))
        return count

    # API de Datos Mejorada #
    def get_contenido_cache(self):
        """Obtiene contenido procesado con cache"""
        cache_key = f'guia_{self.pk}_contenido'
        contenido = cache.get(cache_key)
        if not contenido:
            contenido = self.contenido_procesado
            cache.set(cache_key, contenido, timeout=3600)  # 1 hora de cache
        return contenido

    def get_preguntas_por_componente(self, use_cache=True):
        """Devuelve preguntas agrupadas por componente"""
        if use_cache:
            contenido = self.get_contenido_cache()
        else:
            contenido = self.contenido_procesado
            
        return {
            comp['componente_a_evaluar']: [
                pregunta for bloque in comp['bloques'] 
                for pregunta in bloque['preguntas']
            ]
            for comp in contenido.get('tablas_cuestionario', [])
        }

    def get_estadisticas_componentes(self):
        """Devuelve estadísticas resumidas por componente"""
        stats = {}
        contenido = self.get_contenido_cache()
        
        for componente in contenido.get('tablas_cuestionario', []):
            nombre = componente['componente_a_evaluar']
            total_preguntas = 0
            preguntas_ejemplo = []
            
            for bloque in componente.get('bloques', []):
                preguntas = bloque.get('preguntas', [])
                total_preguntas += len(preguntas)
                if len(preguntas_ejemplo) < 3 and preguntas:
                    preguntas_ejemplo.append(preguntas[0]['texto'][:50] + '...')
            
            stats[nombre] = {
                'total_preguntas': total_preguntas,
                'bloques': len(componente.get('bloques', [])),
                'preguntas_ejemplo': preguntas_ejemplo
            }
        
        return stats

    def generar_resumen_evaluacion(self, usuario_id):
        """Genera resumen de evaluación para un usuario específico"""
        from .models import RespuestaGuia  # Importación local para evitar circular
        
        respuestas = RespuestaGuia.objects.filter(
            guia=self,
            usuario_id=usuario_id
        ).select_related('guia')
        
        total_preguntas = self.total_preguntas
        respondidas = respuestas.count()
        
        # Agrupación por componente
        por_componente = {}
        preguntas_por_componente = self.get_preguntas_por_componente()
        
        for componente, preguntas in preguntas_por_componente.items():
            ids_preguntas = [p['numero_pregunta'] for p in preguntas]
            respuestas_componente = respuestas.filter(
                numero_pregunta__in=ids_preguntas
            ).count()
            
            por_componente[componente] = {
                'total': len(ids_preguntas),
                'respondidas': respuestas_componente,
                'porcentaje': round((respuestas_componente / len(ids_preguntas)) * 100, 2) if ids_preguntas else 0
            }
        
        return {
            'total_preguntas': total_preguntas,
            'respondidas': respondidas,
            'porcentaje': round((respondidas / total_preguntas) * 100, 2) if total_preguntas > 0 else 0,
            'por_componente': por_componente
        }


    def get_absolute_url(self):
        return reverse('guia:detalle', kwargs={'pk': self.pk})

    def _limpiar_texto(self, texto):
        """Limpia el texto eliminando caracteres no deseados y normalizando espacios."""
        if not isinstance(texto, str):
            return str(texto)
        texto = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]', '', texto)
        texto = re.sub(r'\s+', ' ', texto)
        try:
            texto = texto.encode('utf-8').decode('utf-8')
        except UnicodeEncodeError:
            texto = texto.encode('ascii', 'ignore').decode('ascii')
        return texto.strip()

    def _procesar_estructura_para_json(self, estructura):
        """Procesa una estructura (lista o diccionario) para asegurar compatibilidad JSON."""
        if isinstance(estructura, list):
            return [self._procesar_estructura_para_json(item) for item in estructura]
        elif isinstance(estructura, dict):
            return {k: self._procesar_estructura_para_json(v) for k, v in estructura.items()}
        else:
            return self._limpiar_texto(str(estructura))

    def _crear_contenido_procesado_seguro(self, error=None):
        """Helper para crear un diccionario seguro para contenido_procesado."""
        return {'error': error} if error else {}
    
    def _parsear_tabla_docx(self, path):
        doc = Document(path)
        tablas_cuestionario = []
        componente_actual = None
        dentro_tabla_cuestionario = False
        componentes_procesados = set()
        bloque = False

        # Inicializa un contador para las preguntas
        contador_preguntas = 0

        for tabla in doc.tables:
            for fila in tabla.rows:
                celdas = [self._limpiar_texto(cell.text) if cell.text else "" for cell in fila.cells]
                #Descomentar para debug---> print(f"Celdas: {celdas}")  # Impresión de depuración

                if not any(celdas):
                    #Descomentar para debug---> print("Fila vacía, continuando...")
                    continue

                if any("ASPECTOS A VERIFICAR" in c.upper() for c in celdas if c):
                    dentro_tabla_cuestionario = True
                    #Descomentar para debug---> print("Entrando a tabla de cuestionario")
                    continue

                if not dentro_tabla_cuestionario:
                    #Descomentar para debug---> print("Fuera de tabla de cuestionario, continuando...")
                    continue

                if celdas and "Elaborado y aprobado" in celdas[0]:
                    #Descomentar para debug---> print("Tabla 'Elaborado y aprobado' encontrada, terminando...")
                    break

                if len(celdas) >= 3 and len(set(c for c in celdas if c)) == 1:
                    componente_texto = self._limpiar_texto(celdas[0])
                    if componente_texto and componente_texto not in componentes_procesados:
                        componente_actual = {
                            "componente_a_evaluar": componente_texto,
                            "bloques": []
                        }
                        tablas_cuestionario.append(componente_actual)
                        componentes_procesados.add(componente_texto)
                        bloque = None
                        #Descomentar para debug---> print(f"Componente detectado: {componente_texto}")
                        continue

                # Detectar encabezado
                encabezado_detectado = False
                for celda in celdas:
                    if celda and re.search(r":\s*$", celda):
                        encabezado_texto = self._limpiar_texto(celda)
                        if componente_actual:
                            bloque = {
                                "encabezado": encabezado_texto,
                                "preguntas": []
                            }
                            componente_actual["bloques"].append(bloque)
                            #Descomentar para debug--> print(f"Encabezado detectado y añadido: {encabezado_texto}")
                        encabezado_detectado = True
                        break  # Solo un encabezado por fila

                # Detectar preguntas si no hubo encabezado en la misma fila
                if not encabezado_detectado and len(celdas) > 1:
                    if re.match(r"^\d+\.?\s*.+", celdas[1]):
                        numero_pregunta = int(re.search(r"^\d+", celdas[1]).group())
                        texto_pregunta = re.sub(r"^\d+\.?\s*", "", celdas[1]).strip()
                    else:
                        numero_pregunta = None
                        texto_pregunta = celdas[1].strip()

                    if not bloque and componente_actual:
                        bloque = {
                            "encabezado": "",
                            "preguntas": []
                        }
                        componente_actual["bloques"].append(bloque)

                    # Añadir la pregunta al bloque actual con un identificador único
                    if componente_actual and bloque:
                        contador_preguntas += 1
                        bloque["preguntas"].append({
                            "numero_pregunta": contador_preguntas,  # Identificador único
                            "texto": texto_pregunta
                        })
                        #Descomentar para debug--> print(f"[Debug]:preguntas: {bloque}")

        return tablas_cuestionario

    def extraer_contenido_archivo(self):
        """
        Extrae y procesa el contenido del archivo asociado (PDF/DOCX).
        Limpia y estructura componente, propósito y cuestionario.
        """
        if not self.archivo or not self.archivo.archivo:
            raise ValueError("No hay un archivo asociado a esta Guía de Autocontrol.")

        file_path = self.archivo.archivo.path
        mime, _ = mimetypes.guess_type(file_path)
        full_text = ""

        try:
            if mime == 'application/pdf' or file_path.lower().endswith('.pdf'):
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    full_text = "\n".join(page.extract_text() for page in reader.pages if page.extract_text())
                tablas_cuestionario = self._parsear_tabla_pdf(file_path)
            elif mime == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or file_path.lower().endswith('.docx'):
                doc = Document(file_path)
                full_text = "\n".join(para.text for para in doc.paragraphs)
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                        if row_text:
                            full_text += row_text + "\n"
                tablas_cuestionario = self._parsear_tabla_docx(file_path)
            else:
                raise ValueError("Tipo de archivo no soportado para extracción automática.")

            comp_match = re.search(r'COMPONENTE\s+«?([A-ZÁÉÍÓÚÜÑ\s]+)»?', full_text, re.IGNORECASE)
            self.componente = comp_match.group(1).strip() if comp_match else ""

            prop_match = re.search(r'Propósito:\s*([\s\S]*?)(?=Principales fuentes de información para el autocontrol|$)', full_text, re.IGNORECASE)
            if prop_match:
                proposito = prop_match.group(1).strip()
                proposito = re.sub(r'(?<![.!?])\n(?!\n)', ' ', proposito)
                proposito = re.sub(r'\s{2,}', ' ', proposito)
                proposito = re.sub(r'\n\s*\n', '\n', proposito)
                proposito = re.sub(r'(\s*\.\s*)+', '. ', proposito)
                proposito = proposito.replace('•', '')
                lineas = [l.strip() for l in proposito.split('\n') if l.strip() and not re.fullmatch(r'\d+', l.strip())]
                buffer = ''
                parrafos = []
                for line in lineas:
                    line = re.sub(r'^[0-9]+[\s\.-·]*', '', line)
                    if re.match(r'^[A-ZÁÉÍÓÚÜÑ\s]+:$', line):
                        continue
                    elif re.match(r'^(-|\*|•|·|[a-zA-Z]\))', line) or line.startswith(''):
                        continue
                    else:
                        if buffer:
                            buffer += ' '
                        buffer += line
                        if line.endswith('.'):
                            parrafos.append(buffer.strip())
                            buffer = ''
                if buffer:
                    parrafos.append(buffer.strip())
                self.proposito = " ".join(parrafos)

            self.contenido_procesado = {
                "componente": self.componente,
                "proposito": self.proposito,
                "tablas_cuestionario": tablas_cuestionario if tablas_cuestionario else []
            }
            self.save()

        except Exception as e:
            logger.error(f"Error al extraer contenido del archivo {self.archivo.nombre}: {e}")
            raise
        
class RespuestaGuia(models.Model):
    """
    Modelo mejorado de Respuesta con optimizaciones
    """
    RESPUESTA_CHOICES = [
        ('si', 'Sí'),
        ('no', 'No'),
        ('na', 'No Aplica'),
        ('', 'Sin Responder'),
    ]
    
    guia = models.ForeignKey(GuiaAutocontrol, on_delete=models.CASCADE)
    usuario = models.ForeignKey(User, on_delete=models.CASCADE)
    numero_pregunta = models.IntegerField()
    respuesta = models.CharField(
        max_length=4, 
        choices=RESPUESTA_CHOICES, 
        blank=True, 
        null=True, 
        default=''
    )
    fundamentacion = models.TextField(blank=True)
    fecha_respuesta = models.DateTimeField(auto_now=True)
    fecha_modificacion = models.DateTimeField(auto_now=True)
    evidencias = models.ManyToManyField(Archivo, blank=True)
    
    class Meta:
        unique_together = ('guia', 'usuario', 'numero_pregunta')
        verbose_name = 'Respuesta de Guía'
        ordering = ['guia_id', 'numero_pregunta']
        indexes = [
            Index(fields=['guia', 'usuario']),
            Index(fields=['numero_pregunta']),
            Index(fields=['respuesta']),
        ]

    def __str__(self):
        return f"Resp. a P{self.numero_pregunta} ({self.guia}) por {self.usuario}"


class EvaluacionGuia(models.Model):
    """
    Modelo mejorado de Evaluación con optimizaciones
    """
    ESTADO_CHOICES = [
        ('no_iniciada', 'No Iniciada'),
        ('en_progreso', 'En Progreso'),
        ('completada', 'Completada'),
        ('revisada', 'Revisada'),
    ]
    
    usuario = models.ForeignKey(User, on_delete=models.CASCADE, related_name='evaluaciones_guias')
    guia = models.ForeignKey(GuiaAutocontrol, on_delete=models.CASCADE, related_name='evaluaciones')
    fecha_inicio = models.DateTimeField(auto_now_add=True)
    fecha_completado = models.DateTimeField(null=True, blank=True)
    estado = models.CharField(max_length=20, choices=ESTADO_CHOICES, default='no_iniciada')
    porcentaje_cumplimiento = models.DecimalField(
        max_digits=5, 
        decimal_places=2, 
        default=0.00,
        validators=[MinValueValidator(0), MaxValueValidator(100)]
    )
    comentarios = models.TextField(blank=True)
    respuestas_json = models.JSONField(default=dict, blank=True, help_text="Tabla de respuestas agrupadas por componente y pregunta")
    
    # Campos denormalizados para optimización
    total_respuestas = models.PositiveIntegerField(default=0)
    respuestas_si = models.PositiveIntegerField(default=0)
    respuestas_no = models.PositiveIntegerField(default=0)
    respuestas_na = models.PositiveIntegerField(default=0)
    
    class Meta:
        unique_together = ('guia', 'usuario')
        verbose_name = 'Evaluación de Guía'
        ordering = ['-fecha_inicio']
        indexes = [
            Index(fields=['estado']),
            Index(fields=['porcentaje_cumplimiento']),
            Index(fields=['guia', 'usuario']),
        ]

    def __str__(self):
        return f"Evaluación de {self.guia.titulo_guia} por {self.usuario.username}"

    def actualizar_estadisticas(self):
        """
        Actualiza campos denormalizados basado en respuestas
        """
        # Obtener todas las respuestas de este usuario para esta guía
        from .models import RespuestaGuia
        respuestas = RespuestaGuia.objects.filter(guia=self.guia, usuario=self.usuario)
        
        # Calcula y guarda el total de respuestas y el conteo por tipo para este usuario y guía
        self.total_respuestas = respuestas.count()
        self.respuestas_si = respuestas.filter(respuesta='si').count()
        self.respuestas_no = respuestas.filter(respuesta='no').count()
        self.respuestas_na = respuestas.filter(respuesta='na').count()
        
        # Solo contar respuestas válidas (si/no/na) para el porcentaje
        respuestas_validas = self.respuestas_si + self.respuestas_no + self.respuestas_na
        total_preguntas = self.guia.total_preguntas
        
        if total_preguntas > 0:
            self.porcentaje_cumplimiento = round(
                (respuestas_validas  / total_preguntas) * 100, 
                2
            )
        
        if self.porcentaje_cumplimiento >= 100:
            self.estado = 'completada'
            # Solo asigna la fecha si aún no está puesta
            if not self.fecha_completado:
                from django.utils import timezone
                self.fecha_completado = timezone.now()
        else:
            self.estado = 'en_progreso'
            # Si se vuelve a menos de 100%, borra la fecha de completado
            self.fecha_completado = None

        self.save()

    def actualizar_respuestas_json(self):
        """
        Actualiza el campo respuestas_json con la estructura agrupada por componente, número de pregunta, respuesta y fundamentación.
        """
        from .models import RespuestaGuia
        # Obtener todas las respuestas de este usuario para esta guía
        respuestas = RespuestaGuia.objects.filter(guia=self.guia, usuario=self.usuario)
        # Obtener la estructura de componentes y preguntas
        contenido = self.guia.get_contenido_cache()
        tabla = []
        for componente in contenido.get('tablas_cuestionario', []):
            comp_nombre = componente.get('componente_a_evaluar', '')
            comp_respuestas = []
            for bloque in componente.get('bloques', []):
                for pregunta in bloque.get('preguntas', []):
                    num = pregunta.get('numero_pregunta')
                    r = respuestas.filter(numero_pregunta=num).first()
                    comp_respuestas.append({
                        'numero_pregunta': num,
                        'respuesta': r.respuesta if r else '',
                        'fundamentacion': r.fundamentacion if r else ''
                    })
            tabla.append({
                'componente_a_evaluar': comp_nombre,
                'respuestas_guias': comp_respuestas
            })
        self.respuestas_json = {'tabla_respuestas': tabla}
        self.save(update_fields=['respuestas_json'])

    def generar_informe(self):
        """
        Genera informe detallado de la evaluación
        """
        return {
            'guia': self.guia.titulo_guia,
            'usuario': self.usuario.get_full_name(),
            'estado': self.get_estado_display(),
            'porcentaje': self.porcentaje_cumplimiento,
            'fecha_inicio': self.fecha_inicio,
            'fecha_completado': self.fecha_completado,
            'estadisticas': {
                'total_preguntas': self.guia.total_preguntas,
                'respondidas': self.total_respuestas,
                'si': self.respuestas_si,
                'no': self.respuestas_no,
                'na': self.respuestas_na,
            },
            'por_componente': self.guia.generar_resumen_evaluacion(self.usuario_id)['por_componente']
        }
        return informe
