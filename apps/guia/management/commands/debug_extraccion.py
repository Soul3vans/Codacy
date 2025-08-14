"""
Comando de diagnóstico completo para detectar problemas en la extracción de guías
"""
from django.core.management.base import BaseCommand
from apps.guia.models import GuiaAutocontrol, RespuestaGuia, EvaluacionGuia
from apps.dashboard.models import Archivo
from django.db import connection
import os
import PyPDF2
import docx


class Command(BaseCommand):
    help = 'Diagnóstico completo del sistema de guías'

    def add_arguments(self, parser):
        parser.add_argument('--guia_id', type=int, help='ID específico de guía')
        parser.add_argument('--archivo_id', type=int, help='ID específico de archivo')
        parser.add_argument('--fix', action='store_true', help='Intentar reparar problemas encontrados')
        parser.add_argument('--deep', action='store_true', help='Análisis profundo del contenido')
        parser.add_argument('--clean', action='store_true', help='Limpiar y reprocesar todas las guías')

    def handle(self, *args, **options):
        self.stdout.write(self.style.SUCCESS('=' * 60))
        self.stdout.write(self.style.SUCCESS('DIAGNÓSTICO COMPLETO DEL SISTEMA DE GUÍAS'))
        self.stdout.write(self.style.SUCCESS('=' * 60))
        
        if options['clean']:
            self.limpiar_y_reprocesar()
            return
            
        if options['guia_id']:
            self.diagnosticar_guia(options['guia_id'], options.get('fix', False), options.get('deep', False))
        elif options['archivo_id']:
            self.diagnosticar_archivo(options['archivo_id'], options.get('fix', False))
        else:
            self.diagnostico_general(options.get('fix', False))

    def diagnostico_general(self, fix=False):
        """Diagnóstico general del sistema"""
        self.stdout.write('\n1. VERIFICACIÓN DE BASE DE DATOS')
        self.stdout.write('-' * 40)
        
        # Verificar tablas
        with connection.cursor() as cursor:
            cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
            tablas = cursor.fetchall()
            self.stdout.write(f"Tablas encontradas: {len(tablas)}")
            
            # Verificar tabla de guías
            cursor.execute("SELECT COUNT(*) FROM guia_guiaautocontrol;")
            count_guias = cursor.fetchone()[0]
            self.stdout.write(f"Guías en BD: {count_guias}")
            
            # Verificar tabla de archivos
            cursor.execute("SELECT COUNT(*) FROM dashboard_archivo WHERE es_formulario = 1;")
            count_archivos = cursor.fetchone()[0]
            self.stdout.write(f"Archivos formulario: {count_archivos}")

        self.stdout.write('\n2. ANÁLISIS DE ARCHIVOS FÍSICOS')
        self.stdout.write('-' * 40)
        
        archivos = Archivo.objects.filter(es_formulario=True)
        archivos_ok = 0
        archivos_error = 0
        
        for archivo in archivos:
            status = self.verificar_archivo_fisico(archivo)
            if status['ok']:
                archivos_ok += 1
            else:
                archivos_error += 1
                self.stdout.write(self.style.ERROR(f"ERROR - {archivo.nombre}: {status['error']}"))
        
        self.stdout.write(f"Archivos OK: {archivos_ok}")
        self.stdout.write(f"Archivos con error: {archivos_error}")

        self.stdout.write('\n3. ANÁLISIS DE GUÍAS')
        self.stdout.write('-' * 40)
        
        guias = GuiaAutocontrol.objects.all()
        guias_con_contenido = 0
        guias_sin_contenido = 0
        guias_con_errores = 0
        
        for guia in guias:
            status = self.verificar_guia(guia)
            if status['tiene_contenido']:
                guias_con_contenido += 1
            elif status['tiene_errores']:
                guias_con_errores += 1
            else:
                guias_sin_contenido += 1
        
        self.stdout.write(f"Guías con contenido: {guias_con_contenido}")
        self.stdout.write(f"Guías sin contenido: {guias_sin_contenido}")
        self.stdout.write(f"Guías con errores: {guias_con_errores}")

        if fix and (guias_sin_contenido > 0 or guias_con_errores > 0):
            self.stdout.write('\n4. INTENTANDO REPARACIONES')
            self.stdout.write('-' * 40)
            self.reparar_guias()

    def diagnosticar_guia(self, guia_id, fix=False, deep=False):
        """Diagnóstico específico de una guía"""
        try:
            guia = GuiaAutocontrol.objects.get(pk=guia_id)
            self.stdout.write(f'\nDIAGNÓSTICO DE GUÍA ID: {guia_id}')
            self.stdout.write('=' * 50)
            
            # Información básica
            self.stdout.write(f"Título: {guia.titulo_guia}")
            self.stdout.write(f"Activa: {guia.activa}")
            self.stdout.write(f"Fecha procesamiento: {guia.fecha_procesamiento}")
            
            # Verificar archivo asociado
            if not guia.archivo:
                self.stdout.write(self.style.ERROR("✗ NO HAY ARCHIVO ASOCIADO"))
                return
                
            archivo_status = self.verificar_archivo_fisico(guia.archivo)
            if not archivo_status['ok']:
                self.stdout.write(self.style.ERROR(f"✗ ARCHIVO CON PROBLEMAS: {archivo_status['error']}"))
                if fix:
                    self.stdout.write("Intentando reparar...")
                    # Aquí podrías implementar lógica de reparación
                return
            
            self.stdout.write(self.style.SUCCESS(f"✓ Archivo OK: {archivo_status['info']}"))

            # Mostrar el contenido completo extraído del PDF si --fix
            if fix:
                contenido_pdf = self.leer_archivo_con_debug(guia.archivo)
                self.stdout.write("\n=== TEXTO COMPLETO EXTRAÍDO DEL ARCHIVO ===\n")
                self.stdout.write(contenido_pdf)
                self.stdout.write("\n=== FIN TEXTO ARCHIVO ===\n")
                self.stdout.write(f"Contenido leído: {len(contenido_pdf)} caracteres")
            
            # Verificar contenido procesado
            self.stdout.write(f"\nCONTENIDO PROCESADO:")
            self.stdout.write(f"Existe: {bool(guia.contenido_procesado)}")
            
            if guia.contenido_procesado:
                self.stdout.write(f"Tipo: {type(guia.contenido_procesado)}")
                if isinstance(guia.contenido_procesado, dict):
                    self.stdout.write(f"Keys: {list(guia.contenido_procesado.keys())}")
                    if 'preguntas' in guia.contenido_procesado:
                        total_preguntas = len(guia.contenido_procesado['preguntas'])
                        self.stdout.write(f"Preguntas: {total_preguntas}")
                        if total_preguntas > 0:
                            self.stdout.write("Primeras 3 preguntas extraídas:")
                            for p in guia.contenido_procesado['preguntas'][:3]:
                                self.stdout.write(f"{p['numero']}: {p['pregunta'][:100]}")
                    else:
                        self.stdout.write("Preguntas: 0")
                    if 'error' in guia.contenido_procesado:
                        self.stdout.write(self.style.ERROR(f"Error guardado: {guia.contenido_procesado['error']}"))
                    
                    if deep and 'preguntas':
                        self.stdout.write("\nPRIMERAS 3 PREGUNTAS:")
                        for i, p in enumerate('preguntas'[:3]):
                            self.stdout.write(f"{i+1}. Num: {p.get('numero')}, Cat: {p.get('categoria')}")
                            self.stdout.write(f"   Texto: {p.get('texto', '')[:100]}...")
                
                else:
                    self.stdout.write(self.style.WARNING(f"Contenido no es dict: {guia.contenido_procesado}"))
            else:
                self.stdout.write(self.style.WARNING("Sin contenido procesado"))
                
                if fix:
                    self.stdout.write("\nINTENTANDO PROCESAR...")
                    self.procesar_guia_con_debug(guia)
            
            # Verificar evaluaciones y respuestas
            evaluaciones = EvaluacionGuia.objects.filter(guia=guia).count()
            respuestas = RespuestaGuia.objects.filter(evaluacion__guia=guia).count()
            self.stdout.write(f"\nEVALUACIONES: {evaluaciones}")
            self.stdout.write(f"RESPUESTAS: {respuestas}")
            
        except GuiaAutocontrol.DoesNotExist:
            self.stdout.write(self.style.ERROR(f"Guía ID {guia_id} no encontrada"))

    def diagnosticar_archivo(self, archivo_id, fix=False):
        """Diagnóstico específico de un archivo"""
        try:
            archivo = Archivo.objects.get(pk=archivo_id)
            self.stdout.write(f'\nDIAGNÓSTICO DE ARCHIVO ID: {archivo_id}')
            self.stdout.write('=' * 50)
            
            self.stdout.write(f"Nombre: {archivo.nombre}")
            self.stdout.write(f"Es formulario: {archivo.es_formulario}")
            self.stdout.write(f"Fecha subida: {archivo.fecha_subida}")
            
            # Verificar archivo físico
            status = self.verificar_archivo_fisico(archivo)
            if status['ok']:
                self.stdout.write(self.style.SUCCESS(f"✓ Archivo físico OK: {status['info']}"))
                
                # Intentar leer contenido
                if fix:
                    self.stdout.write("\nINTENTANDO LEER CONTENIDO...")
                    contenido = self.leer_archivo_con_debug(archivo)
                    if contenido:
                        self.stdout.write(f"✓ Contenido leído: {len(contenido)} caracteres")
                        self.stdout.write(f"Muestra: {contenido[:200]}...")
                    else:
                        self.stdout.write(self.style.ERROR("✗ No se pudo leer contenido"))
            else:
                self.stdout.write(self.style.ERROR(f"✗ Archivo físico con problemas: {status['error']}"))
            
            # Verificar si tiene guía asociada
            try:
                guia = archivo.guia_autocontrol
                self.stdout.write(f"\n✓ Tiene guía asociada: ID {guia.pk}")
            except GuiaAutocontrol.DoesNotExist:
                self.stdout.write("\n? No tiene guía asociada")
                if fix:
                    self.stdout.write("Creando guía...")
                    self.crear_guia_desde_archivo(archivo)
                    
        except Archivo.DoesNotExist:
            self.stdout.write(self.style.ERROR(f"Archivo ID {archivo_id} no encontrado"))

    def verificar_archivo_fisico(self, archivo):
        """Verifica que el archivo físico esté accesible"""
        try:
            if not archivo.archivo:
                return {'ok': False, 'error': 'archivo.archivo es None'}
            
            file_path = archivo.archivo.path
            
            if not os.path.exists(file_path):
                return {'ok': False, 'error': f'Archivo no existe: {file_path}'}
            
            if not os.access(file_path, os.R_OK):
                return {'ok': False, 'error': 'Sin permisos de lectura'}
            
            size = os.path.getsize(file_path)
            if size == 0:
                return {'ok': False, 'error': 'Archivo vacío'}
            
            ext = os.path.splitext(file_path)[1].lower()
            if ext not in ['.pdf', '.docx', '.doc']:
                return {'ok': False, 'error': f'Tipo no soportado: {ext}'}
            
            return {
                'ok': True, 
                'info': f'{size} bytes, {ext}',
                'path': file_path,
                'size': size,
                'ext': ext
            }
            
        except Exception as e:
            return {'ok': False, 'error': str(e)}

    def verificar_guia(self, guia):
        """Verifica el estado de una guía"""
        status = {
            'tiene_contenido': False,
            'tiene_errores': False,
            'num_preguntas': 0,
            'archivo_ok': False
        }
        
        # Verificar archivo
        if guia.archivo:
            archivo_status = self.verificar_archivo_fisico(guia.archivo)
            status['archivo_ok'] = archivo_status['ok']
        
        # Verificar contenido
        if guia.contenido_procesado:
            if isinstance(guia.contenido_procesado, dict):
                preguntas = guia.contenido_procesado.get('preguntas', [])
                if preguntas:
                    status['tiene_contenido'] = True
                    status['num_preguntas'] = len(preguntas)
                if 'error' in guia.contenido_procesado:
                    status['tiene_errores'] = True
        
        return status

    def procesar_guia_con_debug(self, guia):
        """Procesa una guía con debug detallado"""
        self.stdout.write("--- PROCESANDO GUÍA CON DEBUG ---")
        
        # Limpiar contenido previo
        guia.contenido_procesado = {}
        guia.save()
        
        # Leer archivo
        contenido = self.leer_archivo_con_debug(guia.archivo)
        if not contenido:
            self.stdout.write(self.style.ERROR("No se pudo leer archivo"))
            return
        # Imprime el texto completo extraído del PDF para depuración
        self.stdout.write("\n=== TEXTO COMPLETO EXTRAÍDO DEL PDF (procesar_guia_con_debug) ===\n")
        self.stdout.write(contenido)
        self.stdout.write("\n=== FIN TEXTO PDF ===\n")
        self.stdout.write(f"Contenido leído: {len(contenido)} caracteres")
        
        # Procesar contenido paso a paso
        try:
            # Extractar título
            titulo = self.extraer_titulo_debug(contenido)
            self.stdout.write(f"Título: '{titulo}'")
            
            # Extraer preguntas
            preguntas = self.extraer_preguntas_debug(contenido)
            self.stdout.write(f"Preguntas encontradas: {len(preguntas)}")
            
            # Guardar resultado
            guia.titulo_guia = titulo
            guia.contenido_procesado = {
                'titulo_guia': titulo,
                'preguntas': preguntas,
                'fecha_procesamiento': guia.fecha_procesamiento.isoformat(),
                'debug_info': {
                    'contenido_length': len(contenido),
                    'metodo_procesamiento': 'debug_manual'
                }
            }
            guia.save()
            # Imprimir tabla_cuestionario o preguntas extraídas
            self.stdout.write("\n=== tabla_cuestionario extraída ===")
            import pprint
            pprint.pprint(guia.contenido_procesado.get('tablas_cuestionario') or guia.contenido_procesado.get('preguntas'))
            self.stdout.write("=== FIN tabla_cuestionario ===\n")
            
            self.stdout.write(self.style.SUCCESS("✓ Guía procesada exitosamente"))
            
        except Exception as e:
            self.stdout.write(self.style.ERROR(f"Error procesando: {e}"))
            import traceback
            traceback.print_exc()

    def leer_archivo_con_debug(self, archivo):
        """Lee un archivo con debug paso a paso"""
        try:
            file_path = archivo.archivo.path
            ext = os.path.splitext(file_path)[1].lower()
            
            self.stdout.write(f"Leyendo archivo: {file_path}")
            self.stdout.write(f"Extensión: {ext}")
            
            if ext == '.pdf':
                return self.leer_pdf_debug(file_path)
            elif ext == '.docx':
                return self.leer_docx_debug(file_path)
            else:
                self.stdout.write(f"Tipo no soportado: {ext}")
                return ""
                
        except Exception as e:
            self.stdout.write(self.style.ERROR(f"Error leyendo archivo: {e}"))
            return ""

    def leer_pdf_debug(self, file_path):
        """Lee PDF con debug e imprime el texto extraído completo, mostrando espacios y saltos de línea explícitos. Elimina saltos de línea salvo si la línea termina en punto."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                self.stdout.write(f"PDF páginas: {num_pages}")
                contenido = ""
                for i, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        # Unir todas las líneas, solo dejar salto si termina en punto
                        lines = page_text.split('\n')
                        new_lines = []
                        buffer = ''
                        for idx, line in enumerate(lines):
                            line = line.strip()
                            if not line:
                                continue
                            if buffer:
                                buffer += ' '
                            buffer += line
                            if line.endswith('.'):
                                new_lines.append(buffer)
                                buffer = ''
                        if buffer:
                            new_lines.append(buffer)
                        page_text = '\n'.join(new_lines)
                        contenido += page_text + "\n"
                        debug_text = page_text.replace(' ', '·').replace('\n', '\\n\n')
                        self.stdout.write(f"Página {i+1}: {len(page_text)} chars\n{debug_text}")
                    except Exception as e:
                        self.stdout.write(f"Error página {i+1}: {e}")
                self.stdout.write("\n=== TEXTO COMPLETO EXTRAÍDO DEL PDF (espacios=·, saltos=\\n) ===\n")
                debug_contenido = contenido.replace(' ', '·').replace('\n', '\\n\n')
                self.stdout.write(debug_contenido)
                self.stdout.write("\n=== FIN TEXTO PDF ===\n")
                return contenido
        except Exception as e:
            self.stdout.write(self.style.ERROR(f"Error PDF: {e}"))
            return ""

    def leer_docx_debug(self, file_path):
        """Lee DOCX con debug, mostrando espacios y saltos de línea explícitos. Elimina saltos de línea salvo si la línea termina en punto."""
        try:
            doc = docx.Document(file_path)
            self.stdout.write(f"DOCX párrafos: {len(doc.paragraphs)}")
            lines = []
            for i, paragraph in enumerate(doc.paragraphs):
                lines.extend(paragraph.text.split('\n'))
            # Unir todas las líneas, solo dejar salto si termina en punto
            new_lines = []
            buffer = ''
            for idx, line in enumerate(lines):
                line = line.strip()
                if not line:
                    continue
                if buffer:
                    buffer += ' '
                buffer += line
                if line.endswith('.'):
                    new_lines.append(buffer)
                    buffer = ''
            if buffer:
                new_lines.append(buffer)
            contenido = '\n'.join(new_lines)
            debug_contenido = contenido.replace(' ', '·').replace('\n', '\\n\n')
            self.stdout.write("\n=== TEXTO COMPLETO EXTRAÍDO DEL DOCX (espacios=·, saltos=\\n) ===\n")
            self.stdout.write(debug_contenido)
            self.stdout.write("\n=== FIN TEXTO DOCX ===\n")
            return contenido
        except Exception as e:
            self.stdout.write(self.style.ERROR(f"Error DOCX: {e}"))
            return ""

    def extraer_titulo_debug(self, contenido):
        """Extrae título con debug"""
        import re
        
        patrones = [
            r'GUÍA DE AUTOCONTROL GENERAL ACTUALIZADA\s*(.*?)\s*COMPONENTE',
            r'GUÍA DE AUTOCONTROL\s*(.*?)\s*COMPONENTE',
            r'GUÍA\s*(.*?)\s*COMPONENTE'
        ]
        
        for patron in patrones:
            match = re.search(patron, contenido, re.IGNORECASE | re.DOTALL)
            if match:
                titulo = match.group(1).strip()
                self.stdout.write(f"Título encontrado con patrón: {patron}")
                return titulo
        
        self.stdout.write("No se encontró título")
        return "Título no encontrado"

    def limpiar_pregunta_duplicada(self, numero, texto):
        import re
        patron = r'^' + re.escape(str(numero)) + r'\s*\.?\s*'
        texto_limpio = re.sub(patron, '', texto, count=1)
        return texto_limpio.strip()

    def extraer_preguntas_debug(self, contenido):
        """Extrae preguntas desde tablas, identificando encabezados y preguntas numeradas"""
        import re
        self.stdout.write("\n=== INICIO DEL TEXTO EXTRAÍDO DEL PDF ===\n")
        self.stdout.write(contenido[:2000])
        self.stdout.write("\n=== FIN DEL TEXTO EXTRAÍDO ===\n")
        patrones_inicio = [
            r'Cuestionario de preguntas:',
            r'CUESTIONARIO DE PREGUNTAS:',
            r'Cuestionario:',
            r'CUESTIONARIO:'
        ]
        inicio = None
        for patron in patrones_inicio:
            match = re.search(patron, contenido, re.IGNORECASE)
            if match:
                inicio = match.end()
                self.stdout.write(f"Cuestionario encontrado con: {patron}")
                break
        if not inicio:
            self.stdout.write("No se encontró inicio de cuestionario")
            return []
        bloque = contenido[inicio:inicio+20000]
        self.stdout.write(f"Sección extraída: {len(bloque)} chars")
        # Eliminar encabezados de columnas
        bloque = re.sub(r'NO\.\s*\(1\)\s*ASPECTOS A VERIFICAR\s*\(2\)\s*SÍ\s*\(3\)\s*NO\s*\(4\)\s*Fundamento\s*\(5\)', '', bloque, flags=re.IGNORECASE)
        # Separar por líneas
        lineas = [l.strip() for l in bloque.split('\n') if l.strip()]
        preguntas = []
        encabezado_actual = None
        for i, linea in enumerate(lineas):
            # Encabezado: línea sin número al inicio y no vacía
            if re.match(r'^[A-Za-zÁÉÍÓÚÜÑáéíóúüñ\(\)\s\-]+$', linea) and not re.match(r'^\d+\.', linea):
                encabezado_actual = linea
                self.stdout.write(f"Encabezado detectado: {encabezado_actual}")
                continue
            # Pregunta numerada
            m = re.match(r'^(\d+)\.\s*(.*)', linea)
            if m:
                numero = m.group(1)
                texto = m.group(2).strip()
                # Unir líneas siguientes que no sean encabezado ni pregunta
                j = i + 1
                while j < len(lineas) and not re.match(r'^(\d+\.|[A-Za-zÁÉÍÓÚÜÑáéíóúüñ\(\)\s\-]+)$', lineas[j]):
                    texto += ' ' + lineas[j]
                    j += 1
                texto_limpio = self.limpiar_pregunta_duplicada(numero, texto)
                pregunta = {
                    'numero': numero,
                    'texto': texto_limpio,
                    'encabezado': encabezado_actual
                }
                preguntas.append(pregunta)
                self.stdout.write(f"{numero}. {texto_limpio}")
        self.stdout.write(f"\nTotal preguntas extraídas: {len(preguntas)}")
        return preguntas

    def crear_guia_desde_archivo(self, archivo):
        """Crea una guía desde un archivo"""
        try:
            guia = GuiaAutocontrol.objects.create(
                archivo=archivo,
                titulo_guia=archivo.nombre,
                activa=True
            )
            self.stdout.write(f"✓ Guía creada: ID {guia.pk}")
            
            # Procesar contenido
            self.procesar_guia_con_debug(guia)
            
        except Exception as e:
            self.stdout.write(self.style.ERROR(f"Error creando guía: {e}"))

    def limpiar_y_reprocesar(self):
        """Limpia y reprocesa todas las guías"""
        self.stdout.write("LIMPIANDO Y REPROCESANDO TODAS LAS GUÍAS")
        self.stdout.write("=" * 50)
        
        # Limpiar contenido de todas las guías
        guias = GuiaAutocontrol.objects.all()
        self.stdout.write(f"Limpiando {guias.count()} guías...")
        
        for guia in guias:
            guia.contenido_procesado = {}
            guia.save()
        
        # Reprocesar
        self.stdout.write("Reprocesando...")
        for i, guia in enumerate(guias, 1):
            self.stdout.write(f"Procesando {i}/{guias.count()}: {guia.titulo_guia}")
            self.procesar_guia_con_debug(guia)

    def reparar_guias(self):
        """Intenta reparar guías problemáticas"""
        guias_problematicas = GuiaAutocontrol.objects.filter(
            contenido_procesado__isnull=True
        ) | GuiaAutocontrol.objects.filter(
            contenido_procesado={}
        )
        
        self.stdout.write(f"Reparando {guias_problematicas.count()} guías...")
        
        for guia in guias_problematicas:
            self.stdout.write(f"Reparando: {guia.titulo_guia}")
            self.procesar_guia_con_debug(guia)