from django.core.management.base import BaseCommand
from apps.guia.models import GuiaAutocontrol, RespuestaGuia, EvaluacionGuia
from django.db import connection
from apps.dashboard.models import Archivo 
import os
import PyPDF2 
import docx 
import logging
logger = logging.getLogger(__name__)

def extraer_incisos_tabla2(ruta_archivo):
    from docx import Document
    import re

    doc = Document(ruta_archivo)
    # Selecciona la tabla 2 (índice 1)
    try:
        tabla = doc.tables[1]
    except IndexError:
        print("No existe la tabla 2 en el documento.")
        return

    # Busca encabezado para confirmar que es la tabla correcta
    encabezado = [cell.text.strip() for cell in tabla.rows[0].cells]
    print(f"Encabezado tabla 2: {encabezado}")
    if not encabezado or not encabezado[0].startswith("NO"):
        print("La tabla 2 no tiene el encabezado esperado.")
        return

    print(f"\n--- Todas las celdas de la tabla 2 ---")
    for fila_idx, fila in enumerate(tabla.rows):
        for col_idx, celda in enumerate(fila.cells):
            texto = celda.text.strip()
            print(f"Fila {fila_idx}, Columna {col_idx}: '{texto}'")
    print(f"\n--- Incisos encontrados en la tabla 2 ---")
    for fila_idx, fila in enumerate(tabla.rows):
        for col_idx, celda in enumerate(fila.cells):
            texto = celda.text.strip()
            # Busca incisos en cada línea del texto de la celda
            for linea in texto.splitlines():
                match = re.match(r'^([a-zA-Z0-9]\))\s*(.*)', linea.strip())
                if match:
                    inciso = match.group(1)
                    pregunta = match.group(2)
                    print(f"Fila {fila_idx}, Columna {col_idx}: Inciso '{inciso}' - Pregunta: '{pregunta}'")

def contar_tablas_docx(ruta_archivo):
    from docx import Document
    doc = Document(ruta_archivo)
    print(f"El documento tiene {len(doc.tables)} tablas.")
    for i, tabla in enumerate(doc.tables):
        encabezado = [cell.text.strip() for cell in tabla.rows[0].cells] if tabla.rows else []
        print(f"Tabla {i+1} encabezado: {encabezado}")

class Command(BaseCommand):
    help = 'Diagnóstico completo del sistema de guías'

    def add_arguments(self, parser):
        parser.add_argument('--guia_id', type=int, help='ID específico de guía')
        parser.add_argument('--archivo_id', type=int, help='ID específico de archivo')
        parser.add_argument('--fix', action='store_true', help='Intentar reparar problemas encontrados')
        parser.add_argument('--deep', action='store_true', help='Análisis profundo del contenido')
        parser.add_argument('--clean', action='store_true', help='Limpiar y reprocesar todas las guías')

    def handle(self, *args, **options):
        self.stdout.write(self.style.SUCCESS('=' * 60))
        self.stdout.write(self.style.SUCCESS('DIAGNÓSTICO COMPLETO DEL SISTEMA DE GUÍAS'))
        self.stdout.write(self.style.SUCCESS('=' * 60))

        # Llama a la función de incisos para la tabla 2
        extraer_incisos_tabla2("media/archivos/1-GA_AMBIENTE_DE_CONTROL_AP_14.5.25.0.docx")
        contar_tablas_docx("media/archivos/1-GA_AMBIENTE_DE_CONTROL_AP_14.5.25.0.docx")

        if options['clean']:
            self.limpiar_y_reprocesar()
            return
        
        if options['guia_id']:
            guia_id = options['guia_id']
            deep_analysis = options['deep']
            fix_problems = options['fix']

            try:
                guia = GuiaAutocontrol.objects.get(pk=guia_id)
                self.stdout.write(f"Procesando guía {guia.titulo_guia} (ID: {guia.pk})...")
                if deep_analysis:
                    self.stdout.write(f"Iniciando análisis profundo para Guía ID: {guia.pk}...")
                    guia.extraer_contenido_archivo() # Esta línea guarda los datos en la DB
                    # ¡¡¡LA LÍNEA CLAVE ES ESTA!!!
                    guia.refresh_from_db() # Recarga el objeto de la base de datos para obtener los cambios guardados.
                    
                    self.stdout.write(f"Resultado en contenido_procesado: {guia.contenido_procesado}")
                    
                    if guia.contenido_procesado.get('tablas_cuestionario'):
                        self.stdout.write(self.style.SUCCESS(f"✓ {len(guia.contenido_procesado['tablas_cuestionario'])} tabla(s) de cuestionario encontrada(s)."))
                        for i, tabla in enumerate(guia.contenido_procesado['tablas_cuestionario']):
                            self.stdout.write(f"  Tabla {i+1}: Categoría '{tabla.get('categoria', 'N/A')}', {len(tabla.get('preguntas', []))} preguntas.")
                    else:
                        self.stdout.write(self.style.WARNING("✗ No se encontraron tablas de cuestionario."))
                # ... (resto de tu lógica del comando) ...

            except GuiaAutocontrol.DoesNotExist:
                self.stdout.write(self.style.ERROR(f"Guía con ID {guia_id} no encontrada."))
            except Exception as e:
                    self.stdout.write(self.style.ERROR(f"Error inesperado al diagnosticar la guía: {e}"))
            return # Exit after processing specific guia_id
            

        # ... rest of your handle method logic for other options ...

    # Make sure the procesar_guia_con_debug method is also defined in this class
    # (It was present in your uploaded debug_extraccion.py)
    def procesar_guia_con_debug(self, guia, deep_analysis=False):
        # ... your existing logic for processing guia.extraer_contenido_archivo()
        # and printing debug info ...
        # Remember to include the detailed print statements in guia.extraer_contenido_archivo()
        self.stdout.write(f"Procesando guía {guia.titulo_guia} (ID: {guia.pk})...")
        try:
            # This calls the method in models.py
            guia.extraer_contenido_archivo()
            self.stdout.write(f"Resultado en contenido_procesado: {guia.contenido_procesado}")
            if deep_analysis and guia.contenido_procesado:
                # Add more detailed prints if deep_analysis is true
                pass
        except Exception as e:
            self.stdout.write(self.style.ERROR(f"Error procesando guía {guia.pk}: {e}"))
    
    def diagnostico_general(self, fix=False):
        """Diagnóstico general del sistema"""
        self.stdout.write('\n1. VERIFICACIÓN DE BASE DE DATOS')
        self.stdout.write('-' * 40)
        
        # Verificar tablas
        with connection.cursor() as cursor:
            cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
            tablas = cursor.fetchall()
            self.stdout.write(f"Tablas encontradas: {len(tablas)}")
            
            # Verificar tabla de guías
            cursor.execute("SELECT COUNT(*) FROM guia_guiaautocontrol;")
            count_guias = cursor.fetchone()[0]
            self.stdout.write(f"Guías en BD: {count_guias}")
            
            # Verificar tabla de archivos
            cursor.execute("SELECT COUNT(*) FROM dashboard_archivo WHERE es_formulario = 1;")
            count_archivos = cursor.fetchone()[0]
            self.stdout.write(f"Archivos formulario: {count_archivos}")

        self.stdout.write('\n2. ANÁLISIS DE ARCHIVOS FÍSICOS')
        self.stdout.write('-' * 40)
        
        archivos = Archivo.objects.filter(es_formulario=True)
        archivos_ok = 0
        archivos_error = 0
        
        for archivo in archivos:
            status = self.verificar_archivo_fisico(archivo)
            if status['ok']:
                archivos_ok += 1
            else:
                archivos_error += 1
                self.stdout.write(self.style.ERROR(f"ERROR - {archivo.nombre}: {status['error']}"))
        
        self.stdout.write(f"Archivos OK: {archivos_ok}")
        self.stdout.write(f"Archivos con error: {archivos_error}")

        self.stdout.write('\n3. ANÁLISIS DE GUÍAS')
        self.stdout.write('-' * 40)
        
        guias = GuiaAutocontrol.objects.all()
        guias_con_contenido = 0
        guias_sin_contenido = 0
        guias_con_errores = 0
        
        for guia in guias:
            status = self.verificar_guia(guia)
            if status['tiene_contenido']:
                guias_con_contenido += 1
            elif status['tiene_errores']:
                guias_con_errores += 1
            else:
                guias_sin_contenido += 1
        
        self.stdout.write(f"Guías con contenido: {guias_con_contenido}")
        self.stdout.write(f"Guías sin contenido: {guias_sin_contenido}")
        self.stdout.write(f"Guías con errores: {guias_con_errores}")

        if fix and (guias_sin_contenido > 0 or guias_con_errores > 0):
            self.stdout.write('\n4. INTENTANDO REPARACIONES')
            self.stdout.write('-' * 40)
            self.reparar_guias()

    def diagnosticar_guia(self, guia_id, fix=False, deep=False):
        """Diagnóstico específico de una guía"""
        try:
            guia = GuiaAutocontrol.objects.get(pk=guia_id)
            self.stdout.write(f'\nDIAGNÓSTICO DE GUÍA ID: {guia_id}')
            self.stdout.write('=' * 50)
            
            # Información básica
            self.stdout.write(f"Título: {guia.titulo_guia}")
            self.stdout.write(f"Activa: {guia.activa}")
            self.stdout.write(f"Fecha procesamiento: {guia.fecha_procesamiento}")
            
            # Verificar archivo asociado
            if not guia.archivo:
                self.stdout.write(self.style.ERROR("✗ NO HAY ARCHIVO ASOCIADO"))
                return
                
            archivo_status = self.verificar_archivo_fisico(guia.archivo)
            if not archivo_status['ok']:
                self.stdout.write(self.style.ERROR(f"✗ ARCHIVO CON PROBLEMAS: {archivo_status['error']}"))
                if fix:
                    self.stdout.write("Intentando reparar...")
                    # Aquí podrías implementar lógica de reparación
                return
            
            self.stdout.write(self.style.SUCCESS(f"✓ Archivo OK: {archivo_status['info']}"))
            
            # Verificar contenido procesado
            self.stdout.write(f"\nCONTENIDO PROCESADO:")
            self.stdout.write(f"Existe: {bool(guia.contenido_procesado)}")
            
            if guia.contenido_procesado:
                self.stdout.write(f"Tipo: {type(guia.contenido_procesado)}")
                if isinstance(guia.contenido_procesado, dict):
                    self.stdout.write(f"Keys: {list(guia.contenido_procesado.keys())}")
                    
                    preguntas = guia.contenido_procesado.get('preguntas', [])
                    self.stdout.write(f"Preguntas: {len(preguntas)}")
                    
                    if 'error' in guia.contenido_procesado:
                        self.stdout.write(self.style.ERROR(f"Error guardado: {guia.contenido_procesado['error']}"))
                    
                    if deep and preguntas:
                        self.stdout.write("\nPRIMERAS 3 PREGUNTAS:")
                        for i, p in enumerate(preguntas[:3]):
                            self.stdout.write(f"{i+1}. Num: {p.get('numero')}, Cat: {p.get('categoria')}")
                            self.stdout.write(f"   Texto: {p.get('texto', '')[:100]}...")
                
                else:
                    self.stdout.write(self.style.WARNING(f"Contenido no es dict: {guia.contenido_procesado}"))
            else:
                self.stdout.write(self.style.WARNING("Sin contenido procesado"))
                
                if fix:
                    self.stdout.write("\nINTENTANDO PROCESAR...")
                    self.procesar_guia_con_debug(guia)
            
            # Verificar evaluaciones y respuestas
            evaluaciones = EvaluacionGuia.objects.filter(guia=guia).count()
            respuestas = RespuestaGuia.objects.filter(guia=guia).count()
            
            self.stdout.write(f"\nEVALUACIONES: {evaluaciones}")
            self.stdout.write(f"RESPUESTAS: {respuestas}")
            
        except GuiaAutocontrol.DoesNotExist:
            self.stdout.write(self.style.ERROR(f"Guía ID {guia_id} no encontrada"))

    def diagnosticar_archivo(self, archivo_id, fix=False):
        """Diagnóstico específico de un archivo"""
        try:
            archivo = Archivo.objects.get(pk=archivo_id)
            self.stdout.write(f'\nDIAGNÓSTICO DE ARCHIVO ID: {archivo_id}')
            self.stdout.write('=' * 50)
            
            self.stdout.write(f"Nombre: {archivo.nombre}")
            self.stdout.write(f"Es formulario: {archivo.es_formulario}")
            self.stdout.write(f"Fecha subida: {archivo.fecha_subida}")
            
            # Verificar archivo físico
            status = self.verificar_archivo_fisico(archivo)
            if status['ok']:
                self.stdout.write(self.style.SUCCESS(f"✓ Archivo físico OK: {status['info']}"))
                
                # Intentar leer contenido
                if fix:
                    self.stdout.write("\nINTENTANDO LEER CONTENIDO...")
                    contenido = self.leer_archivo_con_debug(archivo)
                    if contenido:
                        self.stdout.write(f"✓ Contenido leído: {len(contenido)} caracteres")
                        self.stdout.write(f"Muestra: {contenido[:200]}...")
                    else:
                        self.stdout.write(self.style.ERROR("✗ No se pudo leer contenido"))
            else:
                self.stdout.write(self.style.ERROR(f"✗ Archivo físico con problemas: {status['error']}"))
            
            # Verificar si tiene guía asociada
            try:
                guia = archivo.guia_autocontrol
                self.stdout.write(f"\n✓ Tiene guía asociada: ID {guia.pk}")
            except GuiaAutocontrol.DoesNotExist:
                self.stdout.write("\n? No tiene guía asociada")
                if fix:
                    self.stdout.write("Creando guía...")
                    self.crear_guia_desde_archivo(archivo)
                    
        except Archivo.DoesNotExist:
            self.stdout.write(self.style.ERROR(f"Archivo ID {archivo_id} no encontrado"))

    def verificar_archivo_fisico(self, archivo):
        """Verifica que el archivo físico esté accesible"""
        try:
            if not archivo.archivo:
                return {'ok': False, 'error': 'archivo.archivo es None'}
            
            file_path = archivo.archivo.path
            
            if not os.path.exists(file_path):
                return {'ok': False, 'error': f'Archivo no existe: {file_path}'}
            
            if not os.access(file_path, os.R_OK):
                return {'ok': False, 'error': 'Sin permisos de lectura'}
            
            size = os.path.getsize(file_path)
            if size == 0:
                return {'ok': False, 'error': 'Archivo vacío'}
            
            ext = os.path.splitext(file_path)[1].lower()
            if ext not in ['.pdf', '.docx', '.doc']:
                return {'ok': False, 'error': f'Tipo no soportado: {ext}'}
            
            return {
                'ok': True, 
                'info': f'{size} bytes, {ext}',
                'path': file_path,
                'size': size,
                'ext': ext
            }
            
        except Exception as e:
            return {'ok': False, 'error': str(e)}

    def verificar_guia(self, guia):
        """Verifica el estado de una guía"""
        status = {
            'tiene_contenido': False,
            'tiene_errores': False,
            'num_preguntas': 0,
            'archivo_ok': False
        }
        
        # Verificar archivo
        if guia.archivo:
            archivo_status = self.verificar_archivo_fisico(guia.archivo)
            status['archivo_ok'] = archivo_status['ok']
        
        # Verificar contenido
        if guia.contenido_procesado:
            if isinstance(guia.contenido_procesado, dict):
                if 'error' in guia.contenido_procesado:
                    status['tiene_errores'] = True
                else:
                    preguntas = guia.contenido_procesado.get('preguntas', [])
                    if preguntas:
                        status['tiene_contenido'] = True
                        status['num_preguntas'] = len(preguntas)
        
        return status

    def leer_archivo_con_debug(self, archivo):
        """Lee un archivo con debug paso a paso"""
        try:
            file_path = archivo.archivo.path
            ext = os.path.splitext(file_path)[1].lower()
            
            self.stdout.write(f"Leyendo archivo: {file_path}")
            self.stdout.write(f"Extensión: {ext}")
            
            if ext == '.pdf':
                return self.leer_pdf_debug(file_path)
            elif ext == '.docx':
                return self.leer_docx_debug(file_path)
            else:
                self.stdout.write(f"Tipo no soportado: {ext}")
                return ""
                
        except Exception as e:
            self.stdout.write(self.style.ERROR(f"Error leyendo archivo: {e}"))
            return ""

    def leer_pdf_debug(self, file_path):
        """Lee PDF con debug"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                self.stdout.write(f"PDF páginas: {num_pages}")
                
                contenido = ""
                for i, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        contenido += page_text + "\n"
                        self.stdout.write(f"Página {i+1}: {len(page_text)} chars")
                    except Exception as e:
                        self.stdout.write(f"Error página {i+1}: {e}")
                
                return contenido
                
        except Exception as e:
            self.stdout.write(self.style.ERROR(f"Error PDF: {e}"))
            return ""

    def leer_docx_debug(self, file_path):
        """Lee DOCX con debug"""
        try:
            doc = docx.Document(file_path)
            self.stdout.write(f"DOCX párrafos: {len(doc.paragraphs)}")
            
            contenido = ""
            for i, paragraph in enumerate(doc.paragraphs):
                contenido += paragraph.text + "\n"
                if i < 5:  # Mostrar primeros 5
                    self.stdout.write(f"Párrafo {i+1}: {paragraph.text[:100]}...")
            
            return contenido
            
        except Exception as e:
            self.stdout.write(self.style.ERROR(f"Error DOCX: {e}"))
            return ""

    def extraer_titulo_debug(self, contenido):
        """Extrae título con debug"""
        import re
        
        patrones = [
            r'GUÍA DE AUTOCONTROL GENERAL ACTUALIZADA\s*(.*?)\s*COMPONENTE',
            r'GUÍA DE AUTOCONTROL\s*(.*?)\s*COMPONENTE',
            r'GUÍA\s*(.*?)\s*COMPONENTE'
        ]
        
        for patron in patrones:
            match = re.search(patron, contenido, re.IGNORECASE | re.DOTALL)
            if match:
                titulo = match.group(1).strip()
                self.stdout.write(f"Título encontrado con patrón: {patron}")
                return titulo
        
        self.stdout.write("No se encontró título")
        return "Título no encontrado"

    def extraer_preguntas_debug(self, contenido):
        """Extrae preguntas con debug"""
        import re
        
        # Buscar sección de cuestionario
        patrones_inicio = [
            r'Cuestionario de preguntas:',
            r'CUESTIONARIO DE PREGUNTAS:',
            r'Cuestionario:',
            r'CUESTIONARIO:'
        ]
        
        inicio = None
        for patron in patrones_inicio:
            match = re.search(patron, contenido, re.IGNORECASE)
            if match:
                inicio = match.end()
                self.stdout.write(f"Cuestionario encontrado con: {patron}")
                break
        
        if not inicio:
            self.stdout.write("No se encontró inicio de cuestionario")
            return []
        
        # Extraer sección
        seccion = contenido[inicio:inicio+10000]  # Primeros 10k chars
        self.stdout.write(f"Sección extraída: {len(seccion)} chars")
        
        # Buscar preguntas numeradas
        preguntas = []
        patron_pregunta = r'(\d+)\.?\s+([^\d\n]{20,}?)(?=\n\s*\d+\.|\n\n|\Z)'
        
        matches = re.finditer(patron_pregunta, seccion, re.MULTILINE | re.DOTALL)
        
        for match in matches:
            numero = int(match.group(1))
            texto = match.group(2).strip()
            texto = re.sub(r'\s+', ' ', texto)
            
            pregunta = {
                'numero': numero,
                'texto': texto,
                'categoria': 'General'
            }
            preguntas.append(pregunta)
            
            if len(preguntas) <= 3:  # Debug primeras 3
                self.stdout.write(f"P{numero}: {texto[:50]}...")
        
        return preguntas

    def extraer_incisos_tabla2(self, ruta_archivo):
        """Extrae y muestra incisos de la tabla 2 de un archivo DOCX"""
        from docx import Document
        import re

        doc = Document(ruta_archivo)
        # Selecciona la tabla 2 (índice 1)
        try:
            tabla = doc.tables[1]
        except IndexError:
            self.stdout.write("No existe la tabla 2 en el documento.")
            return

        # Busca encabezado para confirmar que es la tabla correcta
        encabezado = [cell.text.strip() for cell in tabla.rows[0].cells]
        self.stdout.write(f"Encabezado tabla 2: {encabezado}")
        if not encabezado or not encabezado[0].startswith("NO"):
            self.stdout.write("La tabla 2 no tiene el encabezado esperado.")
            return

        self.stdout.write(f"\n--- Todas las celdas de la tabla 2 ---")
        for fila_idx, fila in enumerate(tabla.rows):
            for col_idx, celda in enumerate(fila.cells):
                texto = celda.text.strip()
                self.stdout.write(f"Fila {fila_idx}, Columna {col_idx}: '{texto}'")
        self.stdout.write(f"\n--- Incisos encontrados en la tabla 2 ---")
        for fila_idx, fila in enumerate(tabla.rows):
            for col_idx, celda in enumerate(fila.cells):
                texto = celda.text.strip()
                # Busca incisos en cada línea del texto de la celda
                for linea in texto.splitlines():
                    match = re.match(r'^([a-zA-Z0-9]\))\s*(.*)', linea.strip())
                    if match:
                        inciso = match.group(1)
                        pregunta = match.group(2)
                        self.stdout.write(f"Fila {fila_idx}, Columna {col_idx}: Inciso '{inciso}' - Pregunta: '{pregunta}'")

    def crear_guia_desde_archivo(self, archivo):
        """Crea una guía desde un archivo"""
        try:
            guia = GuiaAutocontrol.objects.create(
                archivo=archivo,
                titulo_guia=archivo.nombre,
                activa=True
            )
            self.stdout.write(f"✓ Guía creada: ID {guia.pk}")
            
            # Procesar contenido
            self.procesar_guia_con_debug(guia)
            
        except Exception as e:
            self.stdout.write(self.style.ERROR(f"Error creando guía: {e}"))

    def limpiar_y_reprocesar(self):
        """Limpia y reprocesa todas las guías"""
        self.stdout.write("LIMPIANDO Y REPROCESANDO TODAS LAS GUÍAS")
        self.stdout.write("=" * 50)
        
        # Limpiar contenido de todas las guías
        guias = GuiaAutocontrol.objects.all()
        self.stdout.write(f"Limpiando {guias.count()} guías...")
        
        for guia in guias:
            guia.contenido_procesado = {}
            guia.save()
        
        # Reprocesar
        self.stdout.write("Reprocesando...")
        for i, guia in enumerate(guias, 1):
            self.stdout.write(f"Procesando {i}/{guias.count()}: {guia.titulo_guia}")
            self.procesar_guia_con_debug(guia)

    def reparar_guias(self):
        """Intenta reparar guías problemáticas"""
        guias_problematicas = GuiaAutocontrol.objects.filter(
            contenido_procesado__isnull=True
        ) | GuiaAutocontrol.objects.filter(
            contenido_procesado={}
        )
        
        self.stdout.write(f"Reparando {guias_problematicas.count()} guías...")
        
        for guia in guias_problematicas:
            self.stdout.write(f"Reparando: {guia.titulo_guia}")
            self.procesar_guia_con_debug(guia)